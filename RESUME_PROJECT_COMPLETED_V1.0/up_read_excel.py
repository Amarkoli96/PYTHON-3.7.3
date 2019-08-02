import sys
import xlrd
from xlrd import open_workbook
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document('input.docx')

wb = open_workbook('Excel_Input.xlsm')
#sheetname = input("give employee id for generating resume")
sheetname = sys.argv[1]
sheet=wb.sheet_by_name(sheetname)
number_of_rows = sheet.nrows
number_of_columns = sheet.ncols

values = []

# to find number of projects
project_count=0
for row in range(number_of_rows):
    data  = (sheet.cell(row,0).value)
    if data == "PROJECT NAME":
        project_count=project_count+1

for col in range(1, number_of_columns):
    for row in range(number_of_rows):
        data  = (sheet.cell(row,col).value)
        values.append(data)

find_keys=['*name*','*exp*','*tech*']
value_trav = 0
for paragraph in doc.paragraphs:
    if find_keys[value_trav] in paragraph.text:
        paragraph.text=(values[value_trav])
        if find_keys[value_trav] == '*name*':
            paragraph.style = doc.styles['Heading 1']
        else:
            paragraph.style = doc.styles['Normal']
        value_trav=value_trav+1
        if value_trav >2:
            break

j=0
f=0
for single in range(project_count):
    table = doc.add_table(rows = 1, cols = 1)
    table.style='Table Grid'
    heading_cells = table.rows[0].cells
    heading_cells[0].text = (values[f+7])
    f=f+5

    table_list=["Roles & Responsibilities","Technology","Tools","Client"]

    totle_rows=4
    totle_cols=2
    table = doc.add_table(rows = totle_rows, cols = totle_cols)
    table.style='Table Grid'

    for i in range(totle_rows):
        heading_cells = table.rows[i].cells
        heading_cells[0].text = (table_list[i])
        heading_cells[1].text = (values[8+j])
        j=j+1
    j=j+1
    paragraph = doc.add_paragraph("\n")

if sheet.cell(3,1).value == "" or sheet.cell(3,1).value.lower() == "none":
    Nombre_text = ['TRAININGS UNDERGONE','EDUCATIONAL QUALIFICATION']
    k=4
    for x in range(2):
        table = doc.add_table(rows=1, cols=1, style='Table Grid')
        row = table.rows[0]
        shading_elm_l= parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
        row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_l)
        Nombre_text_formatted = row.cells[0].paragraphs[0].add_run(Nombre_text[x])
        Nombre_text_formatted.font.size = Pt(12)
        Nombre_text_formatted.bold = True
        paragraph=doc.add_paragraph("\n").add_run(values[k])
        k=k+1
        paragraph = doc.add_paragraph("\n")

else:
    Nombre_text1 = ['ACHIVEMENTS']
    table = doc.add_table(rows=1, cols=1, style='Table Grid')
    row = table.rows[0]
    shading_elm_l= parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
    row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_l)
    Nombre_text_formatted = row.cells[0].paragraphs[0].add_run(Nombre_text1[0])
    Nombre_text_formatted.font.size = Pt(12)
    Nombre_text_formatted.bold = True
    paragraph=doc.add_paragraph("\n").add_run(values[3])
    paragraph = doc.add_paragraph("\n")

    Nombre_text = ['TRAININGS UNDERGONE','EDUCATIONAL QUALIFICATION']
    k=4
    for x in range(2):
        table = doc.add_table(rows=1, cols=1, style='Table Grid')
        row = table.rows[0]
        shading_elm_l= parse_xml(r'<w:shd {} w:fill="E6E6E6"/>'.format(nsdecls('w')))
        row.cells[0]._tc.get_or_add_tcPr().append(shading_elm_l)
        Nombre_text_formatted = row.cells[0].paragraphs[0].add_run(Nombre_text[x])
        Nombre_text_formatted.font.size = Pt(12)
        Nombre_text_formatted.bold = True
        paragraph=doc.add_paragraph("\n").add_run(values[k])
        k=k+1
        paragraph = doc.add_paragraph("\n")

doc.save(f'{sheetname}.docx')

print(sheet.cell(3,1).value)
